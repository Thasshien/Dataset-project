# %%
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from pymongo import MongoClient
from PIL import Image
import numpy as np
import pytesseract
import fitz
import os
import io
from docx import Document
import regex as re
import ollama
import json
import numpy as np


from datasets import load_dataset
from sklearn.linear_model import LinearRegression
from sklearn.metrics import cohen_kappa_score
from scipy.stats import pearsonr

client = MongoClient("mongodb+srv://daktrboys05_db_user:gdgclubproject@to-do-list.qmqixqe.mongodb.net/")
db = client["tries_db"]
questions_collection = db["questions"]

# %% [markdown]
# FOR UPDATING DB WITH ANSWER KEY

# %%
file_path = "test.pdf"
answer_file_path = "test_answer.pdf"

# %%
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    collected = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            collected.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    collected.append(cell_text)

    # Headers & footers
    for section in doc.sections:
        header = section.header
        footer = section.footer

        for para in header.paragraphs:
            if para.text.strip():
                collected.append(para.text.strip())

        for para in footer.paragraphs:
            if para.text.strip():
                collected.append(para.text.strip())

    # Embedded images → OCR
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_bytes = rel.target_part.blob
            img = Image.open(io.BytesIO(image_bytes))
            ocr_text = pytesseract.image_to_string(
                img,
                lang="eng",
                config="--psm 6"
            )
            if ocr_text.strip():
                collected.append(ocr_text.strip())

    return "\n".join(collected)

# %%
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    texts = []

    for page_num, page in enumerate(doc):
        # Digital text
        page_text = page.get_text().strip()
        if page_text:
            texts.append(page_text)

        # OCR embedded images
        images = page.get_images(full=True)
        for img in images:
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            img_pil = Image.open(io.BytesIO(image_bytes))
            ocr_text = pytesseract.image_to_string(
                img_pil,
                lang="eng",
                config="--psm 6"
            ).strip()

            if ocr_text:
                texts.append(ocr_text)

        # Full-page OCR fallback
        if not page_text and not images:
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img, lang="eng").strip()

            if ocr_text:
                texts.append(ocr_text)

    return texts


# %%
def extract_text_from_image(image_path):
    img = Image.open(image_path)
    return pytesseract.image_to_string(img).strip()

# %% [markdown]
# SPLITTING FOR TEXT FROM ANSWER KEY

# %%
def split_by_questions(text: str) -> list[str]:
    
    pattern = (
        r"(\[\d+\s*marks?\]\s*Question:.*?)(?=\[\d+\s*marks?\]\s*Question:|$)"
    )
    matches = re.findall(pattern, text, flags=re.IGNORECASE | re.DOTALL)

    question_blocks = []
    for idx, block in enumerate(matches, start=1):
        question_blocks.append(
            f"[QUESTION_{idx}_START]\n{block.strip()}"
        )
    return question_blocks

# %% [markdown]
# SPLITTING FOR TEXT FROM ANSWERS BY STUDENTS

# %%
def split_by_answers(text: str) -> list[str]:
    pattern = r"(Question\s+\d+[\s\S]*?)(?=Question\s+\d+|$)"
    matches = re.findall(pattern, text, flags=re.IGNORECASE)

    answer_blocks = []
    for idx, block in enumerate(matches, start=1):
        answer_blocks.append(
            f"[ANSWER_{idx}_START]\n{block.strip()}"
        )

    return answer_blocks

# %% [markdown]
# FUNCTION TO EXTRACT TEXT FROM ANSWER KEY

# %%
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        pdf_texts = extract_text_from_pdf(file_path)
        full_text = "\n".join(pdf_texts)
        return split_by_questions(full_text)
    
    elif ext == ".docx":
        return (list(extract_text_from_docx(file_path)))

    elif ext in [".png", ".jpg", ".jpeg"]:
        return (list(extract_text_from_image(file_path)))

    else:
        raise ValueError("Unsupported file type")

# %% [markdown]
# FUNCTION TO EXTRACT TEXT FROM STUDENT ANSWERS

# %%
def extract_student_text(answer_file_path):
    ext = os.path.splitext(answer_file_path)[1].lower()

    if ext == ".pdf":
        pdf_texts = extract_text_from_pdf(answer_file_path)
        full_text = "\n".join(pdf_texts)
        return split_by_answers(full_text)

    elif ext == ".docx":
        docx_texts = extract_text_from_docx(answer_file_path)
        full_text = "\n".join(docx_texts)
        return split_by_answers(full_text)

    elif ext in [".png", ".jpg", ".jpeg"]:
        img_texts = extract_text_from_image(answer_file_path)
        full_text = "\n".join(img_texts)
        return split_by_answers(full_text)

    else:
        raise ValueError("Unsupported file type")

# %% [markdown]
# CHUNKING THE EXTRACTED TEXT 

# %%
class embeddingManager:
  def __init__(self,model_name : str = "all-MiniLM-L6-v2"):
    #hugging face model for sentence embedding
    self.model_name = model_name
    self.model = None
    self._load_model()

  def _load_model(self):
    try:
      print(f"Loading embedding model: {self.model_name}")
      self.model = SentenceTransformer(self.model_name)
      print(f"Embedding model loaded successfully.Embedding dimension: {self.model.get_sentence_embedding_dimension()}")
    except Exception as e:
      print(f"Error loading embedding model: {e}")

  def generate_embeddings(self,texts:list[str]) -> np.ndarray:#returns numpy array
    if self.model is None:
      self._load_model()
    print(f"Generating embedding for {len(texts)} texts....")
    embeddings = self.model.encode(texts, show_progress_bar = True)
    print("Embedding generated successfully.")
    return embeddings

# %% [markdown]
# EMBEDDING FOR ANSWER KEY

# %%
embedding_manager = embeddingManager()
texts = extract_text(file_path)
print("number of records: ",len(texts))
print("extracted texts:", texts)
for text in texts:
    print("starts:",text)
    print("\n")
embedding_manager.generate_embeddings(texts)

# %% [markdown]
# EMBEDDING FOR ANSWERS FROM STUDENT

# %%
embedding_manager = embeddingManager()

# %% [markdown]
# GRADING TECHNICAL QUESTION

# %%
def semantic_fallback(question, student_answer, embedding_manager) -> int:
    """
    Fallback scoring using semantic similarity
    when rubric config or LLM grading fails.
    """

    if not student_answer.strip():
        return 0

    student_emb = embedding_manager.generate_embeddings([student_answer])
    anchor_emb = embedding_manager.generate_embeddings([question["question_text"]])

    similarity = cosine_similarity(student_emb, anchor_emb)[0][0]

    max_marks = question["max_marks"]

    # Convert similarity → marks
    score = int(similarity * max_marks)

    # Safety floor so good answers never get 0
    if score == 0:
        score = max(1, int(0.3 * max_marks))

    return score

# %%
def keyword_score(student_answer: str, keywords: list[str]) -> float:
    if not keywords:
        return 0.0

    text = student_answer.lower()
    hits = sum(1 for kw in keywords if kw.lower() in text)
    return hits / len(keywords)


# %%
def solution_chunk_score(
    student_answer: str,
    solution_chunks: list[str],
    embedding_manager
) -> float:

    if not solution_chunks:
        return 0.0

    student_emb = embedding_manager.generate_embeddings([student_answer])[0]

    matched = 0
    for chunk in solution_chunks:
        chunk_emb = embedding_manager.generate_embeddings([chunk])[0]
        sim = cosine_similarity([student_emb], [chunk_emb])[0][0]

        if sim >= 0.65:   # semantic threshold
            matched += 1

    return matched / len(solution_chunks)



# %%
def numeric_rule_score(student_answer: str, numeric_rules: dict) -> float:
    if not numeric_rules:
        return 0.0

    expected_numbers = numeric_rules.get("expected_numbers", [])

    if not expected_numbers:
        return 1.0   # nothing expected → full marks

    text = student_answer.lower()
    matched = 0

    for num in expected_numbers:
        if str(num).lower() in text:
            matched += 1

    return matched / len(expected_numbers)


# %%
def semantic_similarity_score(
    student_answer: str,
    model_answer: str,
    embedding_manager
) -> float:

    student_emb = embedding_manager.generate_embeddings([student_answer])
    model_emb = embedding_manager.generate_embeddings([model_answer])

    return cosine_similarity(student_emb, model_emb)[0][0]


# %% [markdown]
# GRADING WITH LLM

# %%
# def grade_technical_question(
#     student_answer: str,
#     technical_config: dict,
#     embedding_manager,
#     max_marks: int
# ) -> int:

#     # 1. Semantic similarity
#     semantic_score = semantic_similarity_score(
#         student_answer,
#         technical_config["model_answer"],
#         embedding_manager
#     )

#     # 2. Solution chunk coverage
#     chunk_score = solution_chunk_score(
#         student_answer,
#         technical_config.get("solution_chunks", []),
#         embedding_manager
#     )

#     # 3. Keyword coverage
#     kw_score = keyword_score(
#         student_answer,
#         technical_config.get("keywords", [])
#     )

#     # 4. Numeric / rule-based score
#     numeric_score = numeric_rule_score(
#         student_answer,
#         technical_config.get("numeric_rules", {})
#     )

#     # Final weighted score
#     final_score = (
#         0.4 * semantic_score +
#         0.3 * chunk_score +
#         0.2 * kw_score +
#         0.1 * numeric_score
#     ) * max_marks

#     return round(final_score)


# %%
def grade_descriptive_question(
    question: dict,
    student_answer: str,
    embedding_manager
) -> int:

    print("Grading descriptive question...")

    max_marks = question["max_marks"]

    descriptive_cfg = question.get("descriptive_config")
    if descriptive_cfg is None:
        print("⚠️ descriptive_config missing → semantic fallback")
        return semantic_fallback(question, student_answer, embedding_manager)

    rubric = descriptive_cfg.get("rubric")
    if not rubric:
        print("⚠️ rubric missing → semantic fallback")
        return semantic_fallback(question, student_answer, embedding_manager)

    if not student_answer.strip():
        return 0

    rubric_prompt = ""
    trait_max_map = {}

    for r in rubric:
        trait = r["trait"]
        trait_marks = round(r["weight"] * max_marks)
        trait_max_map[trait] = trait_marks

        rubric_prompt += f"""
Trait: {trait}
Max Marks: {trait_marks}
Description: {r['description']}
"""

    prompt = f"""
You are an experienced university examiner.

Evaluate the student answer STRICTLY using the rubric below.

Question:
{question['question_text']}

Rubric (use ONLY these traits and max marks):
{rubric_prompt}

Student Answer:
{student_answer}

SCORING RULES:
- Assign INTEGER marks only.
- Score each trait independently.
- Use values from 0 up to Max Marks.
- Partial credit is allowed.
- Do NOT invent traits.
- If the answer meaningfully addresses a trait, award non-zero marks.

OUTPUT FORMAT (STRICT JSON ONLY):
{{
  "scores": {{
    "<trait_name>": <integer_marks>
  }}
}}
"""

    try:
        response = ollama.generate(
            model="llama3:latest",
            prompt=prompt.strip()
        )

        data = json.loads(response["response"])
        scores = data.get("scores", {})

        total = 0
        for trait, max_trait_marks in trait_max_map.items():
            awarded = int(scores.get(trait, 0))
            awarded = max(0, min(awarded, max_trait_marks))
            total += awarded

        return min(total, max_marks)

    except Exception as e:
        print("⚠️ LLM failed:", e)
        return semantic_fallback(question, student_answer, embedding_manager)


# %% [markdown]
# FETCHING FROM BACKEND

# %%
def fetch_question(exam_id: str, question_number: int):
    question = questions_collection.find_one(
        {
            "exam_id": exam_id,
            "question_number": question_number
        },
        {
            "_id": 0
        }
    )
    return question


# %%
exam_id = "CS_ADV_2025"
final_results = []

answer_texts = extract_student_text(answer_file_path)
print("Extracted student answers:", answer_texts)

for i, student_answer in enumerate(answer_texts):
    question_number = i + 1

    question = fetch_question(exam_id, question_number)
    print(
        f"Q{question_number} question_type raw = {repr(question['question_type'])}"
    )
    if question is None:
        continue

    marks = 0
    
    if question["question_type"] == "TECHNICAL":
        tech = question["technical_config"]
        marks = grade_technical_question(
            student_answer=student_answer,
            technical_config=tech,
            embedding_manager=embedding_manager,
            max_marks=question["max_marks"]
    )


    elif question["question_type"] == "DESCRIPTIVE":
        marks = grade_descriptive_question(question, student_answer, embedding_manager)

    final_results.append({
        "question_number": question_number,
        "marks_awarded": marks,
        "max_marks": question["max_marks"]
    })
print(final_results)

# %% [markdown]
# COSINE SIMILARITY FOR STUDENT ANSWER

# %%
def compute_cosine_similarity(
    student_embeddings: np.ndarray,
    reference_embeddings: np.ndarray
) -> np.ndarray:
    return cosine_similarity(student_embeddings, reference_embeddings)

# %% [markdown]
# GRADING ANSWERS

# %%
def grade_answers(
    similarity_matrix: np.ndarray,
    max_marks: int = 10
) -> list[int]:

    scores = []

    for i in range(len(similarity_matrix)):
        sim = similarity_matrix[i][i] 

        if sim >= 0.85:
            marks = max_marks
        elif sim >= 0.70:
            marks = int(0.7 * max_marks)
        elif sim >= 0.50:
            marks = int(0.4 * max_marks)
        else:
            marks = 0

        scores.append(marks)

    return scores


# %% [markdown]
# SAMPLE IP

# %%
# Student answers (from your extractor)
student_answers = extract_student_text(answer_file_path)

# Reference answers (from answer key / MongoDB / file)
reference_answers = [
    "ACID properties ensure atomicity, consistency, isolation, and durability...",
    "SELECT department, AVG(salary) FROM employees GROUP BY department ORDER BY AVG(salary) DESC LIMIT 5;",
    "Supervised learning uses labeled data while unsupervised learning does not...",
    "Use a hash set to find the longest subarray in O(n) time...",
    "Network latency affects throughput and response time...",
    "db.orders.aggregate([...]) groups and filters documents..."
]

embedding_manager = embeddingManager()

# Generate embeddings
student_embeddings = embedding_manager.generate_embeddings(student_answers)
reference_embeddings = embedding_manager.generate_embeddings(reference_answers)

# Compute similarity
similarity_matrix = compute_cosine_similarity(
    student_embeddings,
    reference_embeddings
)

# Grade
scores = grade_answers(similarity_matrix)

print("Similarity matrix:\n", similarity_matrix)
print("Final scores:", scores)

def extract_rag_features(
    student_answer: str,
    model_answer: str,
    technical_config: dict,
    embedding_manager
):
    semantic = semantic_similarity_score(
        student_answer,
        model_answer,
        embedding_manager
    )

    chunk = solution_chunk_score(
        student_answer,
        technical_config.get("solution_chunks", []),
        embedding_manager
    )

    keyword = keyword_score(
        student_answer,
        technical_config.get("keywords", [])
    )

    numeric = numeric_rule_score(
        student_answer,
        technical_config.get("numeric_rules", {})
    )

    return [semantic, chunk, keyword, numeric]
def train_rag_weights(embedding_manager):
    print("Loading Mohler ASAG dataset...")
    dataset = load_dataset("nkazi/MohlerASAG", split="raw_open_ended")
    dataset = dataset.shuffle(seed=42)

    train_size = int(0.7 * len(dataset))
    train_data = dataset[:train_size]
    test_data = dataset[train_size:]

    X_train, y_train = [], []
    X_test, y_test = [], []

    # Mohler does not have chunks / numeric → dummy config
    dummy_config = {
        "solution_chunks": [],
        "keywords": [],
        "numeric_rules": {}
    }

    print("Extracting training features...")
    for s in train_data:
        X_train.append(
            extract_rag_features(
                s["student_answer"],
                s["reference_answer"],
                dummy_config,
                embedding_manager
            )
        )
        y_train.append(s["score"])

    print("Extracting test features...")
    for s in test_data:
        X_test.append(
            extract_rag_features(
                s["student_answer"],
                s["reference_answer"],
                dummy_config,
                embedding_manager
            )
        )
        y_test.append(s["score"])

    X_train = np.array(X_train)
    X_test = np.array(X_test)
    y_train = np.array(y_train)
    y_test = np.array(y_test)

    print("Training regression model to learn weights...")
    reg = LinearRegression(positive=True)
    reg.fit(X_train, y_train)

    print("Evaluating on held-out test set...")
    y_pred = reg.predict(X_test)

    pearson = pearsonr(y_test, y_pred)[0]
    y_pred_round = np.round(y_pred).astype(int)
    qwk = cohen_kappa_score(y_test, y_pred_round, weights="quadratic")

    print("===== EVALUATION RESULTS =====")
    print("Pearson Correlation:", round(pearson, 4))
    print("QWK:", round(qwk, 4))

    w_sem, w_chunk, w_kw, w_num = reg.coef_
    bias = reg.intercept_

    print("===== LEARNED RAG WEIGHTS =====")
    print("Semantic:", round(w_sem, 4))
    print("Chunk:", round(w_chunk, 4))
    print("Keyword:", round(w_kw, 4))
    print("Numeric:", round(w_num, 4))
    print("Bias:", round(bias, 4))

    return {
        "semantic": w_sem,
        "chunk": w_chunk,
        "keyword": w_kw,
        "numeric": w_num,
        "bias": bias
    }
def grade_technical_question(
    student_answer: str,
    technical_config: dict,
    embedding_manager,
    max_marks: int,
    weights: dict
) -> int:

    semantic_score = semantic_similarity_score(
        student_answer,
        technical_config["model_answer"],
        embedding_manager
    )

    chunk_score = solution_chunk_score(
        student_answer,
        technical_config.get("solution_chunks", []),
        embedding_manager
    )

    kw_score = keyword_score(
        student_answer,
        technical_config.get("keywords", [])
    )

    numeric_score = numeric_rule_score(
        student_answer,
        technical_config.get("numeric_rules", {})
    )

    final_score = (
        weights["semantic"] * semantic_score +
        weights["chunk"] * chunk_score +
        weights["keyword"] * kw_score +
        weights["numeric"] * numeric_score +
        weights["bias"]
    ) * max_marks

    # fairness floor
    final_score = max(final_score, 0.35 * max_marks)

    return round(final_score)


if __name__ == "__main__":
    embedding_manager = embeddingManager()

    # Run calibration ONCE
    rag_weights = train_rag_weights(embedding_manager)

    # Then grading
    run_exam_grading(rag_weights)
