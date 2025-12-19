# %% [markdown]
# IMPORTS NEEDED FOR THE CODE

# %%
import os
import io
import re
import json
import requests
import tempfile
from datetime import datetime

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from pymongo import MongoClient

from PIL import Image
import pytesseract
import fitz
from docx import Document

import ollama

# %% [markdown]
# MongoDB CONNECTION

# %%
client = MongoClient(
    "mongodb+srv://daktrboys05_db_user:gdgclubproject@to-do-list.qmqixqe.mongodb.net/"
)
db = client["tries_db"]
assingment = db["submissions"]
questions_collection = db["questions"]



# %% [markdown]
# CONNECT API

# %%
app = FastAPI()

# %% [markdown]
# REQUEST MODEL

# %%
class EvaluateRequest(BaseModel):
    exam_id: str
    submission_id: str
    teacher_id: str

# %% [markdown]
# SAFE JSON

# %%
def safe_json(text: str):
    try:
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if not match:
            return {}

        raw = match.group(0)

        try:
            return json.loads(raw)
        except Exception:
            pass

        cleaned = raw.replace("'", '"').replace("\n", " ").replace("\t", " ")
        try:
            return json.loads(cleaned)
        except Exception:
            return {}

    except Exception:
        return {}


# %% [markdown]
# DOWNLOAD FILE

# %%
def download_file_from_cloudinary(file_url: str) -> str:
    response = requests.get(file_url, stream=True)
    response.raise_for_status()

    suffix = os.path.splitext(file_url)[1]  # .pdf / .docx / .jpg

    temp_file = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=suffix
    )

    with temp_file as f:
        for chunk in response.iter_content(chunk_size=8192):
            f.write(chunk)

    return temp_file.name
# %% [markdown]
# PDF EXTRACTION

# %%
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    texts = []

    for page in doc:
        page_text = page.get_text().strip()
        if page_text:
            texts.append(page_text)

        for img in page.get_images(full=True):
            base = doc.extract_image(img[0])
            img_pil = Image.open(io.BytesIO(base["image"]))
            ocr = pytesseract.image_to_string(img_pil, lang="eng", config="--psm 6")
            if ocr.strip():
                texts.append(ocr.strip())

    return texts


# %% [markdown]
# DOCX EXTRACTION

# %%
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    collected = []

    for para in doc.paragraphs:
        if para.text.strip():
            collected.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    collected.append(cell.text.strip())

    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                collected.append(para.text.strip())
        for para in section.footer.paragraphs:
            if para.text.strip():
                collected.append(para.text.strip())

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img = Image.open(io.BytesIO(rel.target_part.blob))
            text = pytesseract.image_to_string(img, lang="eng", config="--psm 6")
            if text.strip():
                collected.append(text.strip())

    return "\n".join(collected)


# %% [markdown]
# IMAGE EXTRACTION

# %%
def extract_text_from_image(image_path):
    img = Image.open(image_path)
    return pytesseract.image_to_string(img).strip()


# %% [markdown]
# QUESTION SPLITTER

# %%
def split_by_questions(text: str):
    pattern = r"(\[\d+\s*marks?\]\s*Question:.*?)(?=\[\d+\s*marks?\]\s*Question:|$)"
    matches = re.findall(pattern, text, flags=re.IGNORECASE | re.DOTALL)

    return [
        f"[QUESTION_{i+1}_START]\n{block.strip()}"
        for i, block in enumerate(matches)
    ]


# %% [markdown]
# FILE TYPE HANDELING

# %%
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        pdf_texts = extract_text_from_pdf(file_path)
        full_text = "\n".join(pdf_texts)
        return split_by_questions(full_text)

    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
        return split_by_questions(text)

    elif ext in [".png", ".jpg", ".jpeg"]:
        text = extract_text_from_image(file_path)
        return split_by_questions(text)

    else:
        raise ValueError("Unsupported file type")


# %% [markdown]
# FETCH ASSINGMENT FROM MONGO DB

# %%
def fetch_submission(submission_id: str):
    return assingment.find_one(
        {"_id": submission_id},
        {"answer_file_url": 1}
    )

def get_latest_answerkey_url(teacher_id: str):
    teacher = db.teachers.find_one(
        {"_id": teacher_id},
        {"assignments.answerKey": 1}
    )

    if not teacher or "assignments" not in teacher:
        return None

    latest_url = None
    latest_time = None

    for assignment in teacher["assignments"]:
        answer_key = assignment.get("answerKey")
        if not answer_key:
            continue

        uploaded_at = answer_key.get("uploadedAt")
        if not uploaded_at:
            continue

        if latest_time is None or uploaded_at > latest_time:
            latest_time = uploaded_at
            latest_url = answer_key.get("url")

    return latest_url

# %% [markdown]
# MAIN ENDPOINT

# %%
@app.post("/ingest/anserkey-file")
def ingest_from_file(payload: EvaluateRequest):

    local_file_path = None

    try:
        # 🔥 GET LATEST ANSWER KEY URL BY TIME
        file_url = get_latest_answerkey_url(payload.teacher_id)
        if not file_url:
            raise HTTPException(status_code=404, detail="No answer key found")

        # 🔥 DOWNLOAD AS TEMP FILE
        local_file_path = download_file_from_cloudinary(file_url)

        # 🔥 USE YOUR EXISTING EXTRACTOR
        question_blocks = extract_text(local_file_path)

        inserted_ids = []

        for idx, block in enumerate(question_blocks, start=1):
            doc = {
                "exam_id": payload.exam_id,
                "block_number": idx,
                "raw_text": block,
                "created_at": datetime.utcnow()
            }

            result = questions_collection.insert_one(doc)
            inserted_ids.append(str(result.inserted_id))

        return {
            "exam_id": payload.exam_id,
            "inserted_blocks": len(inserted_ids),
            "ids": inserted_ids
        }

    finally:
        if local_file_path and os.path.exists(local_file_path):
            os.remove(local_file_path)



