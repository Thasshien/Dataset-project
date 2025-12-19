# %% [markdown]
# IMPORTS NEEDED

# %%
import os
import io
import re
import json
import tempfile
import requests

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from pymongo import MongoClient

from PIL import Image
import pytesseract
import fitz
from docx import Document

import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from datetime import datetime

import ollama

# %% [markdown]
# DB CREDENTIALS

# %%
app = FastAPI()

client = MongoClient("mongodb+srv://daktrboys05_db_user:gdgclubproject@to-do-list.qmqixqe.mongodb.net/")
db = client["tries_db"]

questions_collection = db["questions"]
submissions_collection = db["submissions"]
assingment = db["assingment_submission"]


# %% [markdown]
# REQUEST BODY MODEL

# %%
class EvaluateRequest(BaseModel):
    submission_id: str

# %% [markdown]
# EMBEDDING MANAGER

# %%
class embeddingManager:
    def __init__(self):
        self.model = SentenceTransformer("all-MiniLM-L6-v2")

    def generate_embeddings(self, texts):
        return self.model.encode(texts)
embedding_manager = embeddingManager()

# %% [markdown]
# FETCH SUBMISSION

# %%
def fetch_submission(submission_id: str):
    return submissions_collection.find_one(
        {"_id": submission_id},
        {"answer_file_url": 1}
    )

# %% [markdown]
# DOWNLOADING FILE FORM CLOUDINARY

# %%
def download_file_from_cloudinary(file_url: str) -> str:
    response = requests.get(file_url, stream=True)
    response.raise_for_status()

    suffix = os.path.splitext(file_url)[1]  # .pdf / .docx / .jpg

    temp_file = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=suffix
    )

    with temp_file as f:
        for chunk in response.iter_content(chunk_size=8192):
            f.write(chunk)

    return temp_file.name


# %% [markdown]
# EXTRACTING DOCX FILE

# %%
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    collected = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            collected.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    collected.append(cell_text)

    # Headers & footers
    for section in doc.sections:
        header = section.header
        footer = section.footer

        for para in header.paragraphs:
            if para.text.strip():
                collected.append(para.text.strip())

        for para in footer.paragraphs:
            if para.text.strip():
                collected.append(para.text.strip())

    # Embedded images → OCR
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_bytes = rel.target_part.blob
            img = Image.open(io.BytesIO(image_bytes))
            ocr_text = pytesseract.image_to_string(
                img,
                lang="eng",
                config="--psm 6"
            )
            if ocr_text.strip():
                collected.append(ocr_text.strip())

    return "\n".join(collected)

# %% [markdown]
# EXTRACTING PDF FILE

# %%
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    texts = []

    for page_num, page in enumerate(doc):
        # Digital text
        page_text = page.get_text().strip()
        if page_text:
            texts.append(page_text)

        # OCR embedded images
        images = page.get_images(full=True)
        for img in images:
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            img_pil = Image.open(io.BytesIO(image_bytes))
            ocr_text = pytesseract.image_to_string(
                img_pil,
                lang="eng",
                config="--psm 6"
            ).strip()

            if ocr_text:
                texts.append(ocr_text)

        # Full-page OCR fallback
        if not page_text and not images:
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img, lang="eng").strip()

            if ocr_text:
                texts.append(ocr_text)

    return texts


# %% [markdown]
# EXTRACT IMAGE UPLOAD

# %%
def extract_text_from_image(image_path):
    img = Image.open(image_path)
    return pytesseract.image_to_string(img).strip()

# %% [markdown]
# SPLITTING FOR TEXT ANSWERS BY STUDENTS

# %%
def split_by_answers(text: str) -> list[str]:
    pattern = r"(Question\s+\d+[\s\S]*?)(?=Question\s+\d+|$)"
    matches = re.findall(pattern, text, flags=re.IGNORECASE)

    answer_blocks = []
    for idx, block in enumerate(matches, start=1):
        answer_blocks.append(
            f"[ANSWER_{idx}_START]\n{block.strip()}"
        )

    return answer_blocks

# %% [markdown]
# FUNCTION TO EXTRACT FILE TYPE FROM STUDENT ANSWERS

# %%
def extract_student_text(answer_file_path):
    ext = os.path.splitext(answer_file_path)[1].lower()

    if ext == ".pdf":
        pdf_texts = extract_text_from_pdf(answer_file_path)
        full_text = "\n".join(pdf_texts)
        return split_by_answers(full_text)

    elif ext == ".docx":
        docx_texts = extract_text_from_docx(answer_file_path)
        full_text = "\n".join(docx_texts)
        return split_by_answers(full_text)

    elif ext in [".png", ".jpg", ".jpeg"]:
        img_texts = extract_text_from_image(answer_file_path)
        full_text = "\n".join(img_texts)
        return split_by_answers(full_text)

    else:
        raise ValueError("Unsupported file type")

# %% [markdown]
# FETCH QUESTION

# %%
def fetch_question(exam_id: str, question_number: int):
    return questions_collection.find_one(
        {
            "exam_id": exam_id,
            "question_number": question_number
        },
        {"_id": 0}
    )

# %% [markdown]
# SEMANTIC FALLBACK

# %%
def semantic_fallback(question, student_answer, embedding_manager) -> int:
    """
    Fallback scoring using semantic similarity
    when rubric config or LLM grading fails.
    """

    if not student_answer.strip():
        return 0

    student_emb = embedding_manager.generate_embeddings([student_answer])
    anchor_emb = embedding_manager.generate_embeddings([question["question_text"]])

    similarity = cosine_similarity(student_emb, anchor_emb)[0][0]

    max_marks = question["max_marks"]

    # Convert similarity → marks
    score = int(similarity * max_marks)

    # Safety floor so good answers never get 0
    if score == 0:
        score = max(1, int(0.3 * max_marks))

    return score

# %% [markdown]
# KEYWORD SIMILARITY CHECK

# %%
def keyword_score(student_answer: str, keywords: list[str]) -> float:
    if not keywords:
        return 0.0

    text = student_answer.lower()
    hits = sum(1 for kw in keywords if kw.lower() in text)
    return hits / len(keywords)

# %% [markdown]
# ANSWER CHUNK SIMILARITY CHECKER

# %%
def solution_chunk_score(
    student_answer: str,
    solution_chunks: list[str],
    embedding_manager
) -> float:

    if not solution_chunks:
        return 0.0

    student_emb = embedding_manager.generate_embeddings([student_answer])[0]

    matched = 0
    for chunk in solution_chunks:
        chunk_emb = embedding_manager.generate_embeddings([chunk])[0]
        sim = cosine_similarity([student_emb], [chunk_emb])[0][0]

        if sim >= 0.65:   # semantic threshold
            matched += 1

    return matched / len(solution_chunks)



# %% [markdown]
# NUMERIC RULE SCORE

# %%
def numeric_rule_score(student_answer: str, numeric_rules: dict) -> float:
    if not numeric_rules:
        return 0.0

    expected_numbers = numeric_rules.get("expected_numbers", [])

    if not expected_numbers:
        return 1.0   # nothing expected → full marks

    text = student_answer.lower()
    matched = 0

    for num in expected_numbers:
        if str(num).lower() in text:
            matched += 1

    return matched / len(expected_numbers)


# %% [markdown]
# SEMANTIC SIMILARITY SCORE

# %%
def semantic_similarity_score(
    student_answer: str,
    model_answer: str,
    embedding_manager
) -> float:

    student_emb = embedding_manager.generate_embeddings([student_answer])
    model_emb = embedding_manager.generate_embeddings([model_answer])

    return cosine_similarity(student_emb, model_emb)[0][0]


# %% [markdown]
# GRADING WITH LLM

# %%
def grade_technical_question(
    student_answer: str,
    technical_config: dict,
    embedding_manager,
    max_marks: int
) -> dict:

    # ---------- NUMERIC GRADING (UNCHANGED) ----------
    semantic_score = semantic_similarity_score(
        student_answer,
        technical_config["model_answer"],
        embedding_manager
    )

    chunk_score = solution_chunk_score(
        student_answer,
        technical_config.get("solution_chunks", []),
        embedding_manager
    )

    kw_score = keyword_score(
        student_answer,
        technical_config.get("keywords", [])
    )

    numeric_score = numeric_rule_score(
        student_answer,
        technical_config.get("numeric_rules", {})
    )

    marks = round((
        0.4 * semantic_score +
        0.3 * chunk_score +
        0.2 * kw_score +
        0.1 * numeric_score
    ) * max_marks)

    # ---------- LLM FEEDBACK ----------
    prompt = f"""
You are a technical examiner.

Question:
{technical_config.get("question_text", "")}

Model / Expected Approach:
{technical_config["model_answer"]}

Student Answer:
{student_answer}

Known Evaluation Signals:
- Important keywords: {technical_config.get("keywords", [])}
- Key solution steps: {technical_config.get("solution_chunks", [])}

TASK:
Provide constructive technical feedback that:
- Mentions what is correct in the student's approach
- Clearly states what is missing or incorrect
- Suggests specific improvements (logic, complexity, explanation)

Do NOT give marks.
Do NOT repeat the model answer.
Keep feedback concise and technical.
"""

    try:
        response = ollama.chat(
            model="llama3:latest",
            messages=[{"role": "user", "content": prompt.strip()}]
        )

        feedback = response["message"]["content"].strip()

    except Exception as e:
        print("⚠️ Technical feedback LLM failed:", e)
        feedback = "Technical feedback could not be generated."

    return {
        "marks": marks,
        "feedback": feedback
    }


# %% [markdown]
# DESCRIPTIVE CHECKER

# %%
def grade_descriptive_question(
    question: dict,
    student_answer: str,
    embedding_manager
) -> dict:

    print("Grading descriptive question...")

    max_marks = question["max_marks"]

    if not student_answer.strip():
        return {
            "marks": 0,
            "feedback": "No answer was provided."
        }

    descriptive_cfg = question.get("descriptive_config")
    rubric = descriptive_cfg.get("rubric") if descriptive_cfg else None

    rubric_prompt = ""
    trait_max_map = {}

    if rubric:
        for r in rubric:
            trait = r["trait"]
            trait_marks = round(r["weight"] * max_marks)
            trait_max_map[trait] = trait_marks

            rubric_prompt += f"""
Trait: {trait}
Max Marks: {trait_marks}
Description: {r['description']}
"""

    prompt = f"""
You are an experienced university examiner.

Question:
{question['question_text']}

Student Answer:
{student_answer}

{"Rubric (if any):" + rubric_prompt if rubric_prompt else ""}

TASKS:
1. If a rubric is provided, assign INTEGER marks per trait.
2. If no rubric is provided, DO NOT invent traits.
3. Provide constructive feedback explaining what was done well and what can be improved.

SCORING RULES:
- Assign INTEGER marks only.
- Partial credit is allowed.
- Be fair and academic.

OUTPUT FORMAT (STRICT JSON ONLY):
{{
  "scores": {{
    "<trait_name_if_any>": <integer_marks>
  }},
  "feedback": "<clear improvement-oriented review>"
}}
"""

    try:
        response = ollama.chat(
            model="llama3:latest",
            messages=[
                {"role": "user", "content": prompt.strip()}
            ],
            format="json"
        )

        data = json.loads(response["message"]["content"])

        scores = data.get("scores", {})
        feedback = data.get("feedback", "")

        if rubric and trait_max_map:
            total = 0
            for trait, max_trait_marks in trait_max_map.items():
                awarded = int(scores.get(trait, 0))
                awarded = max(0, min(awarded, max_trait_marks))
                total += awarded
            marks = min(total, max_marks)
        else:
            marks = semantic_fallback(question, student_answer, embedding_manager)

        return {
            "marks": marks,
            "feedback": feedback
        }

    except Exception as e:
        print("⚠️ LLM failed:", e)
        marks = semantic_fallback(question, student_answer, embedding_manager)
        return {
            "marks": marks,
            "feedback": "Feedback could not be generated due to an evaluation error."
        }

# %% [markdown]
# EVALUATING FROM THE FILE UPLOADED

# %%
def evaluate_student_answers(
    exam_id: str,
    answer_file_path: str,
    embedding_manager
):
    final_results = []

    answer_texts = extract_student_text(answer_file_path)
    print("Extracted student answers:", answer_texts)
    print(f"Total answers extracted: {len(answer_texts)}")

    for i, student_answer in enumerate(answer_texts):
        question_number = i + 1

        question = fetch_question(exam_id, question_number)
        print(
            f"Q{question_number} question_type raw = {repr(question['question_type'])}"
        )

        if question is None:
            continue

        marks = 0

        if question["question_type"] == "TECHNICAL":
            tech = question["technical_config"]
            marks = grade_technical_question(
                student_answer=student_answer,
                technical_config=tech,
                embedding_manager=embedding_manager,
                max_marks=question["max_marks"]
            )

        elif question["question_type"] == "DESCRIPTIVE":
            marks = grade_descriptive_question(
                question,
                student_answer,
                embedding_manager
            )

        final_results.append({
            "question_number": question_number,
            "marks_awarded": marks,
            "max_marks": question["max_marks"]
        })

    return final_results

def build_submission_document(
    submission_id: str,
    student_id: str,
    exam_id: str,
    final_results: list,
    question_lookup: dict
):
    submission_doc = {
        "_id": submission_id,
        "student_id": student_id,
        "exam_id": exam_id,
        "submitted_at": datetime.utcnow(),
        "questions": {}
    }

    for result in final_results:
        qno = result["question_number"]
        qkey = f"Q{qno}"

        question_meta = question_lookup[qno]
        marks_block = result["marks_awarded"]

        if isinstance(marks_block, dict):
            marks = marks_block.get("marks", 0)
            feedback = marks_block.get("feedback", "")
        else:
            marks = marks_block
            feedback = ""

        submission_doc["questions"][qkey] = {
            "question_number": qno,
            "question_type": question_meta["question_type"],
            "question_ref_id": question_meta["_id"],
            "evaluation": {
                "marks": marks,
                "max_marks": result["max_marks"],
                "feedback": feedback
            }
        }

    return submission_doc

def build_question_lookup(exam_id: str):
    lookup = {}

    questions = questions_collection.find(
        {"exam_id": exam_id},
        {
            "_id": 1,
            "question_number": 1,
            "question_type": 1,
            "question_text": 1
        }
    )

    for q in questions:
        lookup[q["question_number"]] = q

    return lookup

# %% [markdown]
# API ENDPOINT

# %%
@app.post("/evaluate")
def evaluate_exam(payload: EvaluateRequest):

    exam_id = "CS_ADV_2025"
    submission_id = payload.submission_id

    try:
        # Fetch submission (Cloudinary URL already in DB)
        submission = fetch_submission(payload.submission_id)
        if not submission:
            raise HTTPException(status_code=404, detail="Submission not found")

        file_url = submission.get("answer_file_url")
        if not file_url:
            raise HTTPException(status_code=400, detail="No file URL in submission")

        # Download student answer file
        local_file_path = download_file_from_cloudinary(file_url)

        final_results = evaluate_student_answers(
            exam_id=exam_id,
            answer_file_path=local_file_path,
            embedding_manager=embedding_manager
        )

        print("Final Results:", final_results)

        question_lookup = build_question_lookup(exam_id)

        submission_doc = build_submission_document(
            submission_id="SER_S123_CS_ADV_2025",
            student_id="S123",
            exam_id=exam_id,
            final_results=final_results,
            question_lookup=question_lookup
        )

        result = assingment.insert_one(submission_doc)

        print("Inserted document ID:", result.inserted_id)

    finally:
        if local_file_path and os.path.exists(local_file_path):
            os.remove(local_file_path)


